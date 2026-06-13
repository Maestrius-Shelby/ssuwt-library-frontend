import logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
import os

# Настроим логирование
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_document(data):
    try:
        # Создаем новый документ
        doc = Document()

        # Добавление футера с номерами страниц
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]

        # Создание элемента для номера страницы
        run = paragraph.add_run()
        field_code = 'PAGE'
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), field_code)
        run._r.append(fldSimple)

        # Форматирование шрифта (по желанию)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
        # Выравнивание по центру
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Установка ориентации страницы на альбомную
        section = doc.sections[0]
        section.page_width = Cm(21.0)  # Ширина A4 в см (альбомная ориентация)
        section.page_height = Cm(29.7)  # Высота A4 в см
        section.orientation = WD_ORIENT.LANDSCAPE

        # Установка отступов: 2 см со всех сторон
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        # Функция для установки шрифта
        def set_font(paragraph, font_name='Times New Roman', font_size=Pt(14), bold=False):
            if not paragraph.runs:  # Проверяем, есть ли runs в абзаце
                run = paragraph.add_run()  # Создаем новый run, если его нет
            else:
                run = paragraph.runs[0]
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size
            run.bold = bold

        # Список опубликованных учебных изданий и научных трудов
        author_paragraph = doc.add_paragraph("Список опубликованных учебных изданий и научных трудов")
        set_font(author_paragraph, bold=False)
        author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

        # Добавление данных в таблицу
        for i, item in enumerate(data):
            # Заголовок работы и год
            paragraph = doc.add_paragraph()
            title_run = paragraph.add_run(item['title'])
            title_run.bold = True  # Выделяем заголовок жирным
            title_run.font.size = Pt(14)

            year_run = paragraph.add_run(f" / {item['publication_year']} г.")
            year_run.font.size = Pt(14)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            
            # Авторы
            authors = [
                f"{author['human']['fio']} ({author['relationship_type']})"
                for author in item['authors']
            ]
            authors_paragraph = doc.add_paragraph(", ".join(authors))
            set_font(authors_paragraph)
            authors_paragraph.paragraph_format.space_before = Pt(0)
            authors_paragraph.paragraph_format.space_after = Pt(0)

            # Информация о публикации
            publication_info = f"{item['publication_type']} / {item['rating']} / {item['theme']} / " \
                                f"Издатель: {item['journal_publisher']} / {item['count_of_pages']} стр. / " \
                                f"{item['department']['institute']} / {item['department']['name']}"
            publication_paragraph = doc.add_paragraph(publication_info)
            set_font(publication_paragraph)
            publication_paragraph.paragraph_format.space_before = Pt(0)
            publication_paragraph.paragraph_format.space_after = Pt(0)

            # Отступ между элементами
            doc.add_paragraph()

        # Путь для сохранения документа
        file_path = "ssuwt/export/export_results/Отчет_по_результатам_поиска.docx"
        
        # Проверка на существование директории
        if not os.path.exists(os.path.dirname(file_path)):
            os.makedirs(os.path.dirname(file_path))
            logger.info("Создана директория для файла.")

        # Сохранение документа
        doc.save(file_path)
        logger.info(f"Документ успешно сохранен по пути: {file_path}")
    except Exception as e:
        logger.error(f"Произошла ошибка при создании документа: {str(e)}")
        logger.error(e, exc_info=True)

