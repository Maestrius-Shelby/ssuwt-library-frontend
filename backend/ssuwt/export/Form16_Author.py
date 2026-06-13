import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL

def create_document(data):
    # Создаем новый документ
    doc = Document()
    
    # Если `data` — строка, преобразуем в словарь
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except json.JSONDecodeError:
            raise ValueError("Ошибка декодирования JSON")

    print("Тип данных после обработки:", type(data))  # Проверяем тип после обработки

    # Извлекаем `author_fio`, если он есть
    author_FIO = data.pop("author_fio", None)

    # Преобразуем ФИО "Фамилия Имя Отчество" → "Фамилия И.О."
    if author_FIO:
        fio_parts = author_FIO.split()
        if len(fio_parts) == 3:
            author_FIO = f"{fio_parts[0]} {fio_parts[1][0]}.{fio_parts[2][0]}."
        elif len(fio_parts) == 2:
            author_FIO = f"{fio_parts[0]} {fio_parts[1][0]}."

    # Извлекаем список публикаций
    data = data.get("results", [])

    # Проверяем, что `data` — список
    if not isinstance(data, list):
        raise TypeError("Ожидался список публикаций, но получен другой тип данных.")
    
    # Добавление футера с номерами страниц
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]

    # Создание элемента для номера страницы
    run = paragraph.add_run()
    field_code = 'PAGE'
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), field_code)
    run._r.append(fldSimple)

    # Форматирование шрифта (по желанию)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    # Выравнивание по центру
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Установка ориентации страницы на альбомную
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # Ширина A4 в см (альбомная ориентация)
    section.page_height = Cm(21.0)  # Высота A4 в см
    section.orientation = WD_ORIENT.LANDSCAPE

    # Установка отступов: 2 см со всех сторон
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Функция для установки шрифта
    def set_font(paragraph, font_name='Times New Roman', font_size=Pt(14), bold=False):
        run = paragraph.runs[0]
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = font_size
        run.bold = bold

    # Создание таблицы для заголовка
    header_table = doc.add_table(rows=2, cols=1)
    header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Первая строка таблицы
    row1 = header_table.cell(0, 0)
    row1.text = "ФЕДЕРАЛЬНОЕ АГЕНТСТВО МОРСКОГО И РЕЧНОГО ТРАНСПОРТА"
    set_font(row1.paragraphs[0], bold=True)
    row1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row1.paragraphs[0].paragraph_format.space_after = Pt(0)
    row1.paragraphs[0].paragraph_format.line_spacing = Pt(12)

    # Добавление двойной линии между строками
    tc_pr = row1._element.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')

    bottom_thin = OxmlElement('w:bottom')
    bottom_thin.set(qn('w:val'), 'thinThickLargeGap')
    bottom_thin.set(qn('w:sz'), '0')
    bottom_thin.set(qn('w:space'), '0')
    bottom_thin.set(qn('w:color'), 'auto')

    borders.append(bottom_thin)
    tc_pr.append(borders)

    # Вторая строка таблицы
    row2 = header_table.cell(1, 0)
    row2.text = (
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ\n"
        "ВЫСШЕГО ОБРАЗОВАНИЯ\n"
        "«СИБИРСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ ВОДНОГО ТРАНСПОРТА»"
    )
    set_font(row2.paragraphs[0], bold=True)
    row2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Удаление границ вокруг таблицы
    for cell in header_table._cells:
        tc_pr = cell._element.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)

    # Список опубликованных учебных изданий и научных трудов 
    author_paragraph = doc.add_paragraph("Список опубликованных учебных изданий и научных трудов")
    set_font(author_paragraph, bold=False)
    author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_paragraph.paragraph_format.space_before = Pt(0)
    author_paragraph.paragraph_format.space_after = Pt(0)
    
    # ФИО
    author_paragraph = doc.add_paragraph(author_FIO)
    set_font(author_paragraph, bold=True)
    author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_paragraph.paragraph_format.space_before = Pt(0)
    author_paragraph.paragraph_format.space_after = Pt(0)
    author_paragraph.paragraph_format.line_spacing = Pt(12)

    # Добавление линии под ФИО
    p = author_paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # Добавление строки под ФИО
    FIO = doc.add_paragraph("(фамилия, имя, отчество)")
    set_font(FIO, bold=False)
    FIO.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    FIO.paragraph_format.space_before = Pt(1)
    FIO.paragraph_format.space_after = Pt(1)
    FIO.paragraph_format.line_spacing = Pt(15)

    # Создание таблицы
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Отключение автоподстройки ширины таблицы
    table.allow_autofit = True
    
    # Установка общей ширины таблицы
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(Cm(25.7).cm * 567)))  # Установка общей ширины таблицы
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    # Установка ширины колонок
    widths = [43000, 270000, 135000, 234000, 36000, 200200]
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tcPr = cell._element.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # Заголовки таблицы
    headers = ["№\nп/п", "Наименование учебных изданий и научных трудов",
            "Форма учеб. изд. и науч. трудов", "Выходные данные", "Объем", "Соавторы"]

    # Создаем первую строку таблицы с заголовками
    def set_cell_border(cell, **kwargs):
        """
        Устанавливаем границу ячейки.
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Проверка существования тега для границ, если не найден - создать
        namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        tcBorders = tcPr.find('.//w:tcBorders', namespaces=namespace)
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # Устанавливаем границы по всем сторонам
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # Проверка существования тега для конкретной границы
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # Применение атрибутов
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_paragraph = hdr_cells[i].paragraphs[0]
        hdr_paragraph.text = header
        set_font(hdr_paragraph, font_size=Pt(12), bold=True)
        # Горизонтальное выравнивание по центру
        hdr_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Вертикальное выравнивание по центру
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Устанавливаем жирные границы для ячеек заголовка
        set_cell_border(
            hdr_cells[i],
            start={"sz": 14, "val": "single", "color": "#000000", "space": "0"},
            top={"sz": 14, "val": "single", "color": "#000000", "space": "0"},
            end={"sz": 14, "val": "single", "color": "#000000", "space": "0"},
            bottom={"sz": 14, "val": "single", "color": "#000000", "space": "0"},
            insideH={"sz": 14, "val": "single", "color": "#000000", "space": "0"},
            insideV={"sz": 14, "val": "single", "color": "#000000", "space": "0"},
        )

    # Номера заголовков (которые идут после обычных заголовков)
    headers_numbers = ["1", "2", "3", "4", "5", "6"]

    # Добавляем вторую строку с номерами
    hdr_cells1 = table.add_row().cells  # Создаем новую строку
    for i, header_number in enumerate(headers_numbers):
        hdr_paragraph1 = hdr_cells1[i].paragraphs[0]
        hdr_paragraph1.text = header_number  # Добавляем номер
        set_font(hdr_paragraph1, font_size=Pt(12))
        hdr_paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # Сортировка данных по типу публикации
    data_sorted = sorted(data, key=lambda x: x['publication_type'])

    # Инициализация переменной для отслеживания текущего типа публикации
    current_publication_type = None

    # Добавление данных в таблицу
    for i, item in enumerate(data_sorted):
        # Если тип публикации изменился, добавляем строку с объединенной ячейкой
        if item['publication_type'] != current_publication_type:
            current_publication_type = item['publication_type']
            
            # Добавление новой строки с объединенной ячейкой для типа публикации
            merged_row = table.add_row().cells
            merged_cell = merged_row[0]
            merged_cell.text = current_publication_type
            merged_cell.merge(merged_row[-1])  # Объединяем первую ячейку с последней (все ячейки в строке)
            
            # Устанавливаем стиль для объединенной ячейки
            merged_paragraph = merged_cell.paragraphs[0]
            merged_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = merged_paragraph.runs[0]
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run.font.bold = True

        # Добавление записи о публикации
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = item['title']
        row_cells[2].text = ' - '
        row_cells[3].text = f"{item['journal_publisher']}, {item['publication_year']} - {item['count_of_pages']} с."
        row_cells[4].text = str(item['count_of_pages'])

        # Фильтруем авторов, исключая `author_FIO`
        filtered_authors = [
            author['human']['fio']
            for author in item['authors']
            if author['human']['fio'] != author_FIO
        ]
        row_cells[5].text = '\n'.join(filtered_authors)

        # Установить выравнивание по центру в 0 и 4 столбцах
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for cell in row_cells:
            set_font(cell.paragraphs[0])

    # Сохранение документа
    
    # Путь для сохранения документа
        file_path = "ssuwt/export/export_results/Форма_16_по_фамилии.docx"
        
        # Сохранение документа
        doc.save(file_path)
